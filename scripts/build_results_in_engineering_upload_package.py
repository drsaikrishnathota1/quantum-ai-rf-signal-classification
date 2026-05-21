from __future__ import annotations

import csv
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission_files" / "results_in_engineering_upload_package"


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def set_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(31, 78, 121)),
        ("Heading 3", 12, RGBColor(31, 78, 121)),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def new_doc(title: str | None = None) -> Document:
    doc = Document()
    set_margins(doc)
    set_styles(doc)
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
    return doc


def add_para(doc: Document, text: str = "", bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_csv_table(doc: Document, csv_path: Path, title: str, max_rows: int | None = None) -> None:
    doc.add_heading(title, level=2)
    with csv_path.open(newline="", encoding="utf-8") as f:
        rows = list(csv.reader(f))
    if max_rows:
        rows = rows[: max_rows + 1]
    if not rows:
        add_para(doc, "No rows available.")
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, rows[0]):
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_para(doc)


def markdown_to_docx(md_path: Path, out_path: Path, title_override: str | None = None) -> None:
    doc = new_doc()
    lines = md_path.read_text(encoding="utf-8").splitlines()
    pending_blank = False
    for raw in lines:
        line = raw.strip()
        if not line:
            pending_blank = True
            continue
        if line.startswith("# "):
            text = title_override or line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.bold = True
            r.font.name = "Times New Roman"
            r.font.size = Pt(16)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            continue
        if len(line) > 3 and line[0].isdigit() and ". " in line[:4]:
            add_number(doc, line.split(". ", 1)[1].strip())
            continue
        if line.startswith("`") and line.endswith("`"):
            p = doc.add_paragraph()
            r = p.add_run(line.strip("`"))
            r.italic = True
            continue
        if pending_blank:
            pending_blank = False
        add_para(doc, line)
    doc.save(out_path)


def build_declaration() -> None:
    doc = new_doc("Declaration of Interests Statement")
    add_para(
        doc,
        "The author declares no known competing financial interests or personal relationships "
        "that could have appeared to influence the work reported in this paper."
    )
    add_para(doc, "Author: Dr. Sai Krishna Thota")
    add_para(doc, "Date: May 21, 2026")
    doc.save(OUT / "01_Declaration_of_Interests_Statement.docx")


def build_highlights() -> None:
    doc = new_doc("Highlights")
    highlights = [
        "Robust RF classifiers are evaluated under jamming and low-SNR stress.",
        "Full RadioML2016.10A GPU validation covers 220,000 examples.",
        "Classical ML, raw-IQ CNN, and quantum-inspired kernels are compared.",
        "Robustness-drop metrics reveal stress-specific model failure modes.",
        "No quantum-advantage claim is made; results support cautious benchmarking.",
    ]
    for item in highlights:
        add_bullet(doc, item)
    doc.save(OUT / "02_Highlights.docx")


def build_title_page() -> None:
    markdown_to_docx(ROOT / "manuscript" / "TITLE_PAGE.md", OUT / "04_Title_Page_Editable.docx")


def build_cover_letter() -> None:
    markdown_to_docx(ROOT / "manuscript" / "COVER_LETTER_DRAFT.md", OUT / "05_Cover_Letter.docx")


def build_manuscript() -> None:
    doc = new_doc()
    md = ROOT / "manuscript" / "MANUSCRIPT_DRAFT.md"
    lines = md.read_text(encoding="utf-8").splitlines()
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:].strip())
            r.bold = True
            r.font.name = "Times New Roman"
            r.font.size = Pt(16)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            add_bullet(doc, line[2:].strip())
        elif len(line) > 3 and line[0].isdigit() and ". " in line[:4]:
            add_number(doc, line.split(". ", 1)[1].strip())
        elif line.startswith("Keywords:"):
            add_para(doc, line)
        elif line.startswith("`") and line.endswith("`"):
            p = doc.add_paragraph()
            r = p.add_run(line.strip("`"))
            r.italic = True
        else:
            add_para(doc, line)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Data Availability", level=1)
    add_para(
        doc,
        "Code, configuration files, result tables, generated figures, and reproducibility instructions "
        "are available at https://github.com/drsaikrishnathota1/quantum-ai-rf-signal-classification. "
        "Synthetic datasets are generated by repository scripts. RadioML data are not redistributed and "
        "must be obtained from the original data provider or validated mirror under applicable license terms."
    )
    doc.add_heading("Funding", level=1)
    add_para(doc, "This research received no external funding.")
    doc.add_heading("Declaration of Interests", level=1)
    add_para(doc, "The author declares no known competing financial interests or personal relationships.")
    doc.add_heading("Author Contribution", level=1)
    add_para(doc, "The author conceived the study, designed the simulation protocol, ran the experiments, interpreted the results, and prepared the manuscript.")
    doc.save(OUT / "03_Manuscript.docx")


def build_tables() -> None:
    doc = new_doc("Editable Tables")
    add_para(doc, "These tables are provided as an editable version for journal upload.")
    table_specs = [
        ("RadioML full benchmark clean performance", "table_radioml2016_full_clean_performance.csv"),
        ("RadioML full benchmark method configuration", "table_radioml2016_full_method_config.csv"),
        ("RadioML full benchmark robustness drops", "table_radioml2016_full_robustness_drop.csv"),
        ("RadioML full benchmark robustness metrics", "table_radioml2016_full_robustness_metrics.csv"),
        ("Synthetic clean performance", "table_synthetic_clean_performance.csv"),
        ("Synthetic accuracy by SNR", "table_synthetic_accuracy_by_snr.csv"),
        ("Synthetic robustness drops", "table_synthetic_robustness_drop.csv"),
        ("Synthetic model robustness summary", "table_synthetic_model_robustness_summary.csv"),
    ]
    for title, filename in table_specs:
        add_csv_table(doc, ROOT / "manuscript_assets" / "tables" / filename, title, max_rows=12)
    doc.save(OUT / "06_Tables_Editable.docx")


def build_supplement() -> None:
    doc = new_doc("Supplementary Material")
    doc.add_heading("Reproducibility Package", level=1)
    add_para(doc, "Repository:")
    add_para(doc, "https://github.com/drsaikrishnathota1/quantum-ai-rf-signal-classification")
    doc.add_heading("Core Commands", level=1)
    commands = [
        ".venv/bin/python scripts/run_synthetic_pipeline.py --samples-per-class 500 --cnn-epochs 18",
        "python scripts/prepare_radioml2016_npz.py --input data/radioml/RML2016.10a_dict_optimized.pkl --out data/radioml/radioml2016_10a_clean.npz",
        "python scripts/add_stress_conditions_to_npz.py --input data/radioml/radioml2016_10a_clean.npz --out data/radioml/radioml2016_10a_stress.npz",
        "python scripts/train_pilot_classifiers.py --data data/radioml/radioml2016_10a_stress.npz --out results/radioml2016_classical --max-train-examples 30000 --max-test-examples 10000",
        "python scripts/train_cnn_iq_baseline.py --data data/radioml/radioml2016_10a_stress.npz --out results/radioml2016_cnn --epochs 40 --batch-size 512 --max-train-examples 80000 --max-test-examples 20000 --device cuda",
        "python scripts/train_quantum_inspired_kernel.py --data data/radioml/radioml2016_10a_stress.npz --out results/radioml2016_quantum_kernel --qubits 5 --max-train-per-class 120 --max-test-per-class 80",
        "python scripts/aggregate_radioml_full_analysis.py",
    ]
    for command in commands:
        add_bullet(doc, command)
    doc.add_heading("Important Scope Note", level=1)
    add_para(
        doc,
        "The package contains synthetic pilot results and full RadioML2016.10A GPU benchmark "
        "results. The RadioML raw dataset and large binary model files are not redistributed."
    )
    doc.save(OUT / "07_Supplementary_Material_Unmarked.docx")


def build_research_data_readme() -> None:
    doc = new_doc("Research Data README")
    add_para(doc, "This file explains the research-data handling for journal upload.")
    doc.add_heading("Synthetic Data", level=1)
    add_para(doc, "Synthetic IQ datasets are reproducible from repository scripts and are not tracked as binary files.")
    doc.add_heading("RadioML Dataset", level=1)
    add_para(
        doc,
        "RadioML2016.10A is used for public-benchmark validation. The raw dataset file is not redistributed "
        "in this repository. Users must obtain the dataset from the original provider or validated public mirror "
        "and comply with applicable license terms. The derived tables and figures are included for reproducibility."
    )
    doc.add_heading("Code Repository", level=1)
    add_para(doc, "https://github.com/drsaikrishnathota1/quantum-ai-rf-signal-classification")
    doc.save(OUT / "08_Research_Data_README.docx")


def build_mapping() -> None:
    mapping = """# Results in Engineering Upload Mapping

Use this mapping for the journal's file-type dropdown.

## Required Starred Items

| Journal dropdown item | Upload this file |
|---|---|
| Declaration of Interests Statement | `01_Declaration_of_Interests_Statement.docx` |
| Highlights | `02_Highlights.docx` |
| Manuscript | `03_Manuscript.docx` |
| Title Page (Editable version) | `04_Title_Page_Editable.docx` |

## Recommended Optional Items

| Journal dropdown item | Upload this file |
|---|---|
| Cover letter | `05_Cover_Letter.docx` |
| Table (Editable version) | `06_Tables_Editable.docx` |
| Figure | `Figure_1_RadioML_Clean_Accuracy.png` |
| Figure | `Figure_2_RadioML_Robustness_Drop.png` |
| Figure | `Figure_3_Synthetic_Clean_Accuracy.png` |
| Figure | `Figure_4_Synthetic_Accuracy_By_SNR.png` |
| Figure | `Figure_5_Synthetic_Robustness_Heatmap.png` |
| Supplementary Material (unmarked version, for production) | `07_Supplementary_Material_Unmarked.docx` |
| Research Data | `08_Research_Data_README.docx` |

## Do Not Upload For First Submission

| Journal dropdown item | Reason |
|---|---|
| Response to Reviews | Only for a revised manuscript after reviewer comments |
| Video | Not needed for this paper |
| Co-submission to Data in Brief | Not needed unless we write a separate data article |
| Co-submission to MethodsX | Not needed unless we write a separate methods article |

## Readiness Note

This package is aligned to the journal upload categories and includes the full GPU-based RadioML2016.10A benchmark results.
"""
    (OUT / "UPLOAD_MAPPING.md").write_text(mapping, encoding="utf-8")


def copy_figures() -> None:
    figure_map = {
        "fig_radioml2016_clean_accuracy.png": "Figure_1_RadioML_Clean_Accuracy.png",
        "fig_radioml2016_robustness_drop.png": "Figure_2_RadioML_Robustness_Drop.png",
        "fig_synthetic_clean_accuracy.png": "Figure_3_Synthetic_Clean_Accuracy.png",
        "fig_synthetic_accuracy_by_snr.png": "Figure_4_Synthetic_Accuracy_By_SNR.png",
        "fig_synthetic_robustness_drop_heatmap.png": "Figure_5_Synthetic_Robustness_Heatmap.png",
    }
    for src, dest in figure_map.items():
        shutil.copy2(ROOT / "manuscript_assets" / "figures" / src, OUT / dest)


def make_zip() -> None:
    zip_path = OUT.parent / "results_in_engineering_upload_package.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUT.iterdir()):
            zf.write(path, arcname=path.name)


def main() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    OUT.mkdir(parents=True, exist_ok=True)
    build_declaration()
    build_highlights()
    build_manuscript()
    build_title_page()
    build_cover_letter()
    build_tables()
    build_supplement()
    build_research_data_readme()
    copy_figures()
    build_mapping()
    make_zip()
    print(f"Wrote upload package to {OUT}")
    print(f"Wrote zip to {OUT.parent / 'results_in_engineering_upload_package.zip'}")


if __name__ == "__main__":
    main()
